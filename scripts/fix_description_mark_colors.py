# -*- coding: utf-8 -*-
"""
修复职业页技能描述内丢失的彩色标识 ●（v2，含牧师神圣领域 docx 兜底）。

数据源（优先级）：
  1. 职业页/数据/*.json 的 field_runs / description_entries / level_upgrades[].line_runs
  2. 牧师神圣领域技能：对应 牧师子分支/神圣领域-*.docx 的 run 颜色

两类修复：
  A. 真丢色：描述/条件/效果/升级文本中的 ● 按源颜色包上 <span style="color:#XXX">。
  B. 伪文本：docx「标识」行被误并入描述/效果，网页出现黑色「标识:●●」文本，
     若卡片已有彩色标识行则删除（HTML + JSON 同步清理）。

用法：
    python scripts/fix_description_mark_colors_v2.py            # 执行修复
    python scripts/fix_description_mark_colors_v2.py --verify   # 仅校验
"""
from __future__ import annotations

import argparse
import html as html_lib
import json
import re
import shutil
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "职业页" / "数据"
PAGES = ROOT / "职业页"
ELECTRON = ROOT / "electron-app" / "职业页"
DEITY_DIR = ROOT / "牧师子分支"

PAGE_JSON = [
    ("吟游诗人", "吟游诗人.json"), ("圣骑士", "圣骑士.json"), ("奇械师", "奇械师.json"),
    ("德鲁伊", "德鲁伊.json"), ("战士", "战士.json"), ("术士", "术士.json"),
    ("武僧", "武僧.json"), ("法师", "法师.json"), ("游荡者", "游荡者.json"),
    ("牧师", "牧师.json"), ("猎人", "猎人.json"), ("萨满祭司", "萨满祭司.json"),
    ("蛮斗士", "蛮斗士.json"), ("魔契师", "魔契师.json"), ("通用天赋树", "通用天赋树.json"),
]

COND_LABEL_RE = re.compile(r"^(前置条件|额外条件)[:：]")
LABEL_PREFIXES = ("前置条件：", "额外条件：", "描述：")
ARTIFACT_RE = re.compile(r"^\s*标识[:：]\s*●+\s*$")

# 已知例外：丰收的喜悦(pr-d-life-2) 的描述 ● 在当前任何 牧师子分支 docx 中均无颜色依据，
# 保留黑色文本，不猜测颜色。
KNOWN_NO_SOURCE = {"pr-d-life-2"}


def norm(t: str) -> str:
    return re.sub(r"\s+", "", t or "")


def esc(t: str) -> str:
    return html_lib.escape(t, quote=False)


def strip_label(text: str) -> str:
    for pre in LABEL_PREFIXES:
        if text.startswith(pre):
            return text[len(pre):]
    if COND_LABEL_RE.match(text):
        i = text.find("：")
        if i < 0:
            i = text.find(":")
        return text[i + 1:]
    return text


def cell_text(raw: str) -> str:
    return re.sub(r"<[^>]+>", "", raw)


def plain_dots(raw: str) -> int:
    no_color = re.sub(r'<span[^>]*style="[^"]*color:[^"]*"[^>]*>.*?</span>', "", raw, flags=re.S)
    return re.sub(r"<[^>]+>", "", no_color).count("●")


def inside_colored_span(html: str, idx: int) -> bool:
    head = html[:idx]
    last_open = head.rfind("<span")
    last_close = head.rfind("</span>")
    return last_open > last_close and "color" in head[last_open:idx]


def apply_segments(inner: str, segments) -> str:
    out = inner
    pos = 0
    for text, color in segments:
        if not text or not color:
            continue
        idx = out.find(text, pos)
        if idx < 0:
            continue
        if inside_colored_span(out, idx):
            pos = idx + len(text)
            continue
        span = f'<span style="color:{color}">{esc(text)}</span>'
        out = out[:idx] + span + out[idx + len(text):]
        pos = idx + len(span)
    return out


def json_segments(runs) -> list:
    return [(r.get("text", ""), r.get("color")) for r in (runs or []) if r.get("color") and r.get("text")]


def docx_skill_rows(docx_path: Path, skill_name: str):
    d = docx.Document(str(docx_path))
    for t in d.tables:
        if not t.rows:
            continue
        if t.rows[0].cells[0].text.strip() != skill_name:
            continue
        rows = []
        for row in t.rows:
            label = row.cells[0].text.strip()
            if label.startswith("标识") or label.startswith("费用"):
                continue
            full = ""
            segs = []
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    t2 = run.text
                    full += t2
                    col = None
                    try:
                        col = run.font.color.rgb if run.font.color and run.font.color.type is not None else None
                    except Exception:
                        col = None
                    if col and t2:
                        segs.append((t2, "#" + str(col)))
            if segs:
                rows.append((full, segs))
        return rows
    return []


def apply_to_cell(art: str, cls: str, match_text: str, segments, report_key, sid, report) -> str:
    if cls == "desc-text":
        pat = r'<span class="desc-text">(.*?)</span>'
    elif cls == "cond-text":
        pat = r'<span class="cond-text">(.*?)</span>'
    elif cls == "effect-cell":
        pat = r'<div class="effect-cell">(.*?)</div>'
    else:
        pat = r'<div class="upgrade-cell">(.*?)</div>'
    want = norm(match_text)
    for m in list(re.finditer(pat, art, re.S)):
        if norm(cell_text(m.group(1))) != want:
            continue
        if plain_dots(m.group(1)) == 0:
            continue
        new_inner = apply_segments(m.group(1), segments)
        if new_inner != m.group(1):
            opening = m.group(0)[:m.group(0).index(">") + 1]
            closing = m.group(0)[m.group(0).rfind("<"):]
            art = art[:m.start()] + opening + new_inner + closing + art[m.end():]
            report[report_key].append(sid)
        break
    return art


def remove_artifact_cells(art: str, sid, report) -> str:
    m = re.search(
        r'<div class="desc-cell"><span class="desc-label">描述：</span><span class="desc-text">(.*?)</span></div>',
        art, re.S)
    if m and ARTIFACT_RE.match(cell_text(m.group(1))) and '标识：</span><span class="attr-val">' in art:
        art = art[:m.start()] + art[m.end():]
        report["B_removed"].append(sid)
    for m in list(re.finditer(r'<div class="effect-cell">(.*?)</div>', art, re.S)):
        if ARTIFACT_RE.match(cell_text(m.group(1))) and '标识：</span><span class="attr-val">' in art:
            art = art[:m.start()] + art[m.end():]
            report["B_removed"].append(sid)
    return art


def fix_html(html: str, skills_by_id: dict, docx_rows_by_name: dict, report: dict) -> str:
    art_re = re.compile(r'<article\b[^>]*\bid="(?P<id>[^"]+)"[^>]*>.*?</article>', re.S)

    def rebuild(match: re.Match) -> str:
        sid = match.group("id")
        art = match.group(0)
        sk = skills_by_id.get(sid)
        if sk is None:
            return art

        fr = sk.get("field_runs") or {}
        for key, runs in fr.items():
            segs = json_segments(runs)
            if not segs:
                continue
            concat = "".join(r.get("text", "") for r in runs)
            if key == "描述":
                art = apply_to_cell(art, "desc-text", strip_label(concat), segs, "A_desc", sid, report)
            elif key in ("前置条件", "额外条件"):
                art = apply_to_cell(art, "cond-text", strip_label(concat), segs, "A_cond", sid, report)

        for entry in sk.get("description_entries") or []:
            segs = json_segments(entry.get("runs") or [])
            if segs:
                art = apply_to_cell(art, "effect-cell", entry.get("text", ""), segs, "A_effect", sid, report)

        for up in sk.get("level_upgrades") or []:
            segs = json_segments(up.get("line_runs") or [])
            if segs:
                full = (up.get("label", "") or "") + (up.get("text", "") or "")
                art = apply_to_cell(art, "upgrade-cell", full, segs, "A_upgrade", sid, report)

        name = sk.get("name")
        for row_text, segs in docx_rows_by_name.get(name, []):
            stripped = strip_label(row_text)
            art = apply_to_cell(art, "desc-text", stripped, segs, "A_desc", sid, report)
            art = apply_to_cell(art, "cond-text", stripped, segs, "A_cond", sid, report)
            art = apply_to_cell(art, "effect-cell", stripped, segs, "A_effect", sid, report)
            art = apply_to_cell(art, "upgrade-cell", row_text, segs, "A_upgrade", sid, report)

        return remove_artifact_cells(art, sid, report)

    return art_re.sub(rebuild, html)


def load_deity_skills():
    """牧师·神圣领域.json 的所有技能（含 deity 与 source_file 映射）。"""
    jp = DATA / "牧师·神圣领域.json"
    data = json.loads(jp.read_text(encoding="utf-8"))
    out = {}
    for dom_name, dom in (data.get("domains") or {}).items():
        src = dom.get("source_file", "")
        for sk in dom.get("skills") or []:
            out[sk["id"]] = {**sk, "_docx": src}
    return out


def build_deity_docx_rows(deity_skills: dict) -> dict:
    """技能名 -> [(row_text, segments)]（从领域 docx 提取）。"""
    cache = {}
    out = {}
    for sk in deity_skills.values():
        src = sk.get("_docx", "")
        if not src:
            continue
        path = DEITY_DIR / src
        if not path.exists():
            continue
        if src not in cache:
            cache[src] = {}
            d = docx.Document(str(path))
            for t in d.tables:
                if not t.rows:
                    continue
                nm = t.rows[0].cells[0].text.strip()
                if not nm:
                    continue
                rows = []
                for row in t.rows:
                    label = row.cells[0].text.strip()
                    if label.startswith("标识") or label.startswith("费用"):
                        continue
                    full = ""
                    segs = []
                    for para in row.cells[0].paragraphs:
                        for run in para.runs:
                            t2 = run.text
                            full += t2
                            col = None
                            try:
                                col = run.font.color.rgb if run.font.color and run.font.color.type is not None else None
                            except Exception:
                                col = None
                            if col and t2:
                                segs.append((t2, "#" + str(col)))
                    if segs:
                        rows.append((full, segs))
                if rows:
                    cache[src][nm] = rows
        out.update(cache[src])
    return out


def fix_page(page: str, json_file: str, report: dict, deity_skills=None, deity_docx=None) -> None:
    hp = PAGES / f"{page}.html"
    jp = DATA / json_file
    data = json.loads(jp.read_text(encoding="utf-8"))
    by_id = {sk["id"]: sk for sk in data["skills"]}
    if deity_skills:
        by_id.update(deity_skills)
    html = hp.read_text(encoding="utf-8")
    html = fix_html(html, by_id, deity_docx or {}, report)
    hp.write_text(html, encoding="utf-8", newline="")
    shutil.copyfile(hp, ELECTRON / f"{page}.html")


def clean_json(json_file: str) -> int:
    jp = DATA / json_file
    data = json.loads(jp.read_text(encoding="utf-8"))
    changed = 0
    for sk in data["skills"]:
        fields = sk.get("fields") or {}
        if ARTIFACT_RE.match(fields.get("描述", "")):
            fields.pop("描述", None)
            changed += 1
        if "description" in sk:
            sk["description"] = [x for x in sk["description"] if not ARTIFACT_RE.match(x)]
        if "description_entries" in sk:
            sk["description_entries"] = [
                e for e in sk["description_entries"] if not ARTIFACT_RE.match(e.get("text", ""))
            ]
    jp.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8", newline="")
    return changed


def scan_plain(path: Path) -> int:
    s = path.read_text(encoding="utf-8")
    no_color = re.sub(r'<span[^>]*style="[^"]*color:[^"]*"[^>]*>.*?</span>', "", s, flags=re.S)
    return re.sub(r"<[^>]+>", "", no_color).count("●")


def verify() -> list:
    errs = []
    warns = []
    for page, _ in PAGE_JSON:
        s = (PAGES / f"{page}.html").read_text(encoding="utf-8")
        for m in re.finditer(r'<article\b[^>]*\bid="([^"]+)"[^>]*>(.*?)</article>', s, re.S):
            sid = m.group(1)
            no_color = re.sub(r'<span[^>]*style="[^"]*color:[^"]*"[^>]*>.*?</span>', "", m.group(2), flags=re.S)
            plain = re.sub(r"<[^>]+>", "", no_color).count("●")
            if plain == 0:
                continue
            if sid in KNOWN_NO_SOURCE:
                warns.append(f"{page}/{sid}: 无原文件颜色依据，保留黑色 ● x{plain}")
                continue
            errs.append(f"{page}/{sid}: 仍有纯文本 ● {plain} 个")
    for w in warns:
        print("  WARN", w)
    return errs


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--verify", action="store_true")
    args = parser.parse_args()

    if args.verify:
        errs = verify()
        if errs:
            print("VERIFY FAIL")
            for e in errs:
                print("  -", e)
            return 1
        print("VERIFY PASS：全部主页面无纯文本 ●")
        return 0

    report = {"A_desc": [], "A_cond": [], "A_effect": [], "A_upgrade": [], "B_removed": []}
    deity_skills = load_deity_skills()
    deity_docx = build_deity_docx_rows(deity_skills)
    for page, jf in PAGE_JSON:
        fix_page(page, jf, report, deity_skills if page == "牧师" else None,
                 deity_docx if page == "牧师" else None)
        n = clean_json(jf)
        shutil.copyfile(DATA / jf, ELECTRON / "数据" / jf)
        print(f"{page}: A描述={len(report['A_desc'])} A条件={len(report['A_cond'])} "
              f"A效果={len(report['A_effect'])} A升级={len(report['A_upgrade'])} "
              f"B删除={len(report['B_removed'])} JSON清理={n}")

    errs = verify()
    if errs:
        print("SYNC FAIL")
        for e in errs:
            print("  -", e)
        return 1
    print("SYNC OK：全部页面已修复并同步 electron-app")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
